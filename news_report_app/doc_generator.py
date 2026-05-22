"""Word document generation with strict Arabic RTL support."""

import os
from pathlib import Path
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from models import NewsItem


class DocExporter:
    """Handles Word document generation with Arabic RTL formatting."""
    
    ARABIC_FONTS = ["Calibri", "Traditional Arabic", "Simplified Arabic", "Arial"]
    
    def __init__(self, template_path: Optional[str] = None):
        """Initialize the exporter with optional template path."""
        self.template_path = template_path
        self.output_dir = Path(__file__).parent / "output"
        self.output_dir.mkdir(exist_ok=True)
    
    def _create_default_template(self) -> Document:
        """Create a default document with Arabic header/footer."""
        doc = Document()
        
        # Set default RTL for the document
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.right_margin = Inches(1)
        section.left_margin = Inches(1)
        
        # Add header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "تقرير الأخبار"
        self._apply_rtl_formatting(header_para)
        
        # Add footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"تم إنشاء التقرير في {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        self._apply_rtl_formatting(footer_para)
        
        return doc
    
    def _apply_rtl_formatting(self, paragraph, font_size=12, is_heading=False) -> None:
        """Apply strict RTL formatting to a paragraph with Calibri font and 1.5 line spacing."""
        paragraph.paragraph_format.bidi = True
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Set 1.5 line spacing
        paragraph.paragraph_format.line_spacing = Pt(18)  # 1.5 * 12pt
        
        for run in paragraph.runs:
            run._element.set(qn('w:rtl'), '1')
            run.font.name = self.ARABIC_FONTS[0]  # Calibri
            run.font.element.rPr.set(qn('w:lang'), 'ar-SA')
            if is_heading:
                run.font.size = Pt(16)
                run.font.bold = True
            else:
                run.font.size = Pt(font_size)
    
    def _add_rtl_run(self, paragraph, text: str, **kwargs) -> None:
        """Add a run with RTL formatting to a paragraph."""
        run = paragraph.add_run(text)
        run._element.set(qn('w:rtl'), '1')
        run.font.name = self.ARABIC_FONTS[0]  # Calibri
        
        # Apply optional formatting
        if 'bold' in kwargs:
            run.bold = kwargs['bold']
        if 'italic' in kwargs:
            run.italic = kwargs['italic']
        if 'size' in kwargs:
            run.font.size = Pt(kwargs['size'])
        if 'color' in kwargs:
            run.font.color.rgb = kwargs['color']
        
        # Set language
        rPr = run._element.get_or_add_rPr()
        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), 'ar-SA')
        rPr.append(lang)
    
    def _add_heading(self, doc: Document, title: str, category: str = "") -> None:
        """Add a news title heading with RTL formatting and large font size."""
        para = doc.add_paragraph()
        
        # Add category badge if exists
        if category:
            cat_run = para.add_run(f"[{category}] ")
            cat_run.font.size = Pt(14)
            cat_run.font.bold = True
            cat_run.font.color.rgb = RGBColor(76, 175, 80)  # Green color
        
        # Add title with larger font
        title_run = para.add_run(title)
        title_run._element.set(qn('w:rtl'), '1')
        title_run.font.name = self.ARABIC_FONTS[0]  # Calibri
        title_run.font.size = Pt(20)  # Large title
        title_run.font.bold = True
        
        # Set language for title
        rPr = title_run._element.get_or_add_rPr()
        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), 'ar-SA')
        rPr.append(lang)
        
        para.paragraph_format.bidi = True
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_after = Pt(12)
    
    def _add_source(self, doc: Document, source: str) -> None:
        """Add source line with RTL formatting."""
        para = doc.add_paragraph()
        self._add_rtl_run(para, f"المصدر: {source}", italic=True, size=10, 
                         color=RGBColor(80, 80, 80))
        para.paragraph_format.space_after = Pt(8)
    
    def _add_content(self, doc: Document, content: str) -> None:
        """Add content paragraph with RTL, justified formatting and 1.5 line spacing."""
        para = doc.add_paragraph()
        self._add_rtl_run(para, content, size=12)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = Pt(18)  # 1.5 line spacing
        para.paragraph_format.space_after = Pt(12)
    
    def _add_image(self, doc: Document, image_path: str) -> bool:
        """Add image with caption, aligned with text (right-aligned for RTL). Returns True if successful."""
        if not image_path:
            return False
        
        # Normalize the path to handle different OS formats
        normalized_path = os.path.normpath(image_path)
        
        if not os.path.exists(normalized_path):
            print(f"Warning: Image not found: {normalized_path} (original: {image_path})")
            print(f"Current working directory: {os.getcwd()}")
            return False
        
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align with RTL text
            
            run = para.add_run()
            run.add_picture(normalized_path, width=Inches(5))
            
            # Add caption
            caption = doc.add_paragraph()
            self._add_rtl_run(caption, "شكل: صورة الخبر", italic=True, size=10)
            caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            caption.paragraph_format.line_spacing = Pt(18)  # 1.5 line spacing
            
            return True
        except Exception as e:
            print(f"Warning: Could not add image {normalized_path}: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _add_coordinates(self, doc: Document, coordinates: str) -> None:
        """Add coordinates paragraph with RTL formatting."""
        para = doc.add_paragraph()
        self._add_rtl_run(para, f"📍 {coordinates}", size=11)
        para.paragraph_format.space_after = Pt(8)
    
    def _add_incident_time(self, doc: Document, incident_time: str) -> None:
        """Add incident time paragraph with RTL formatting."""
        para = doc.add_paragraph()
        self._add_rtl_run(para, f"⏰ وقت الحدث: {incident_time}", size=11)
        para.paragraph_format.space_after = Pt(8)
    
    def _add_recommendation(self, doc: Document, recommendation: str) -> None:
        """Add recommendation with special formatting."""
        para = doc.add_paragraph()
        self._add_rtl_run(para, f"💡 {recommendation}", italic=True, size=11)
        
        # Add subtle shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'F0F0F0')
        para._element.get_or_add_pPr().append(shading)
        
        para.paragraph_format.space_after = Pt(12)
    
    def generate(self, news_items: List[NewsItem]) -> str:
        """Generate Word document from news items grouped by category and return file path."""
        # Load or create document
        if self.template_path and os.path.exists(self.template_path):
            doc = Document(self.template_path)
        else:
            doc = self._create_default_template()
        
        # Group news items by category
        categories_dict = {}
        for item in news_items:
            cat = item.category or "عام"
            if cat not in categories_dict:
                categories_dict[cat] = []
            categories_dict[cat].append(item)
        
        # Define category order (common categories first)
        category_order = ["سياسة", "اقتصاد", "رياضة", "تكنولوجيا", "صحة", "تعليم", "ثقافة", "حوادث", "طقس", "منوعات", "عام"]
        sorted_categories = sorted(categories_dict.keys(), 
                                   key=lambda x: category_order.index(x) if x in category_order else len(category_order))
        
        # Add each category section
        first_item = True
        for category in sorted_categories:
            items_in_category = categories_dict[category]
            
            # Add category heading
            if not first_item:
                doc.add_page_break()
            first_item = False
            
            cat_heading = doc.add_paragraph()
            cat_run = cat_heading.add_run(f"─═ {category} ═─")
            cat_run.font.size = Pt(18)
            cat_run.font.bold = True
            cat_run.font.color.rgb = RGBColor(33, 150, 243)  # Blue color
            cat_heading.paragraph_format.bidi = True
            cat_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cat_heading.paragraph_format.space_before = Pt(20)
            cat_heading.paragraph_format.space_after = Pt(20)
            
            # Add each news item in the category
            for i, item in enumerate(items_in_category):
                # Title with category badge (smaller now since we have category heading)
                self._add_heading(doc, item.title, "")
                
                # Source
                self._add_source(doc, item.source)
                
                # Content
                self._add_content(doc, item.content)
                
                # Image (if exists)
                if item.image_path:
                    self._add_image(doc, item.image_path)
                
                # Coordinates (if exists)
                if item.coordinates:
                    self._add_coordinates(doc, item.coordinates)
                
                # Incident time (if exists)
                if item.incident_time:
                    self._add_incident_time(doc, item.incident_time)
                
                # Recommendation (if exists)
                if item.recommendation:
                    self._add_recommendation(doc, item.recommendation)
                
                # Separator line between items (except last in category)
                if i < len(items_in_category) - 1:
                    separator = doc.add_paragraph()
                    sep_run = separator.add_run("─" * 40)
                    sep_run.font.size = Pt(8)
                    sep_run.font.color.rgb = RGBColor(150, 150, 150)
                    separator.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    separator.paragraph_format.space_before = Pt(10)
                    separator.paragraph_format.space_after = Pt(10)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"تقرير_الأخبار_{timestamp}.docx"
        filepath = self.output_dir / filename
        
        # Save document
        doc.save(str(filepath))
        
        return str(filepath)